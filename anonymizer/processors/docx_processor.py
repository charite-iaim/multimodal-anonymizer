"""
Agentic DOCX file processor using LLM with tool-calling for anonymization.

This processor handles Microsoft Word (.docx) files by:
1. Extracting text content from the DOCX file
2. Using the same agentic approach as the text processor
3. Writing the anonymized content back to a new DOCX file while preserving formatting
"""

import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Set, Tuple
import random

from docx import Document
from docx.shared import Inches, Pt
from langchain_core.messages import HumanMessage, ToolMessage
from pydantic import BaseModel, Field

from ..base_processor import FileProcessor
from ..config import AnonymizerConfig
from ..llm_factory import create_chat_llm
from ..tools.time_shift_tool import shift_datetime
from ..tools.redact_tool import redact_text
from ..retry_utils import retry_with_backoff, RetryConfig, create_retry_callback
from ..prompt_config import PromptConfig, DEFAULT_PROMPT_CONFIG


class DateTimeShift(BaseModel):
    """A date/time shift in the text."""
    original_value: str = Field(description="Original date/time value")
    shifted_value: str = Field(description="Shifted date/time value")
    context: str = Field(description="Brief context where this date was found")


class DocxProcessor(FileProcessor):
    """
    Processor for DOCX files using LLM with tool-calling.

    This processor implements a three-phase approach:
    1. Time-Shift Phase: Regex extraction + shift_datetime tool for all dates
    2. Anonymization Phase: LLM uses redact_text tool for PII
    3. Verification Phase: LLM verifies and fixes any issues
    """

    def __init__(
        self,
        config: AnonymizerConfig,
        time_offset_days: Optional[int] = None,
        prompt_config: Optional[PromptConfig] = None
    ):
        """
        Initialize agentic DOCX processor.

        Args:
            config: Configuration object with LLM settings
            time_offset_days: Fixed offset for time shifting. If None, a random offset is generated.
            prompt_config: Custom prompt configuration. If None, uses default prompts.
        """
        super().__init__(config)

        # Prompt configuration
        self.prompt_config = prompt_config or DEFAULT_PROMPT_CONFIG

        # Generate random offset if not provided
        if time_offset_days is None:
            # Random offset between 1-3 years in days
            self.time_offset_days = random.randint(365, 1095)

            # Random sign (+/-)
            if random.random() < 0.5:
                self.time_offset_days = -self.time_offset_days
        else:
            self.time_offset_days = time_offset_days

        # Configure retry settings for LLM calls
        self.retry_config = RetryConfig(
            max_retries=3,
            initial_delay=2.0,
            max_delay=60.0,
            exponential_base=2.0,
            jitter=True,
        )

        # Initialize LLM with tools for phase 1 (date shifting)
        self.llm_with_tools = create_chat_llm(
            config=config,
            timeout=600,
            max_tokens=16000,
            tools=[shift_datetime],
        )

        # Initialize LLM with tools for phase 2 (PII anonymization)
        self.llm_anonymize = create_chat_llm(
            config=config,
            timeout=600,
            max_tokens=16000,
            tools=[redact_text],
        )

        # Initialize LLM with tools for phase 3 (verification)
        self.llm_verify = create_chat_llm(
            config=config,
            timeout=600,
            max_tokens=16000,
            tools=[shift_datetime, redact_text],
        )

    def can_process(self, file_path: Path) -> bool:
        """Check if file is a DOCX file."""
        return file_path.suffix.lower() == ".docx"

    def extract_content(self, file_path: Path) -> str:
        """Extract text content from DOCX file."""
        doc = Document(file_path)

        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    def anonymize(self, input_path: Path, output_path: Path, verify: bool = True) -> None:
        """
        Anonymize DOCX file using agentic LLM approach.

        Steps:
        1. Read DOCX file and extract text
        2. Phase 1: Extract and shift all dates/times using regex
        3. Phase 2: Use LLM with redact_text tool to anonymize PII
        4. Phase 3: Verification agent checks and fixes any issues
        5. Apply all collected replacements directly to DOCX and save

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save anonymized DOCX file
            verify: Whether to run the verification phase (default: True)
        """
        # Convert to Path if string
        input_path = Path(input_path) if isinstance(input_path, str) else input_path
        output_path = Path(output_path) if isinstance(output_path, str) else output_path

        print(f"Processing DOCX: {input_path.name}")
        print(f"Time offset: {self.time_offset_days} days")

        # Step 1: Read DOCX file
        doc = Document(input_path)
        content = self.extract_content(input_path)
        original_content = content
        print(f"Read {len(content)} characters from DOCX")

        if not content.strip():
            print("Empty DOCX file, saving as-is")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return

        # Collect ALL replacements from all phases
        all_replacements: Dict[str, str] = {}

        # Step 2: Phase 1 - Time shifting with regex extraction
        print("\n=== Phase 1: Time Shifting ===")
        shifted_content, dates_shifted, date_replacements = self._phase1_shift_times(content)
        all_replacements.update(date_replacements)
        print(f"Shifted {len(dates_shifted)} date/time values")

        # Step 3: Phase 2 - Anonymize other PII (agentic with redact_text tool)
        print("\n=== Phase 2: PII Anonymization ===")
        anonymized_content, pii_redactions, pii_replacements = self._phase2_anonymize_pii(shifted_content)
        all_replacements.update(pii_replacements)
        print(f"Applied {pii_redactions} PII redactions")

        # Step 4: Phase 3 - Verification (optional but recommended)
        if verify:
            print("\n=== Phase 3: Iterative Verification ===")
            max_iterations = 3
            total_fixes = 0

            for iteration in range(max_iterations):
                print(f"\n  Verification pass {iteration + 1}/{max_iterations}...")
                anonymized_content, fixes_applied, verify_replacements = self._phase3_verify_and_fix(
                    original_content, anonymized_content
                )
                all_replacements.update(verify_replacements)
                total_fixes += fixes_applied
                print(f"  Pass {iteration + 1}: Applied {fixes_applied} fixes.")

                if fixes_applied == 0:
                    print(f"  No more issues found. Verification complete after {iteration + 1} pass(es).")
                    break
            else:
                print(f"  Warning: Completed {max_iterations} passes with {total_fixes} total fixes.")
                print(f"  Consider reviewing the output manually for any remaining PIIs.")

            print(f"\nVerification summary: Applied {total_fixes} total fixes across all passes.")

        # Step 5: Apply all collected replacements directly to DOCX
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._apply_replacements_to_docx(doc, all_replacements, output_path)
        print(f"Saved anonymized DOCX to: {output_path}")
        print(f"Total replacements applied: {len(all_replacements)}")

        # Save JSON with details (if debug mode)
        if self.config.save_debug_files:
            json_output_path = output_path.with_suffix('.json')
            self._save_json_output(
                dates_shifted,
                pii_redactions,
                input_path,
                output_path,
                json_output_path
            )
            print(f"Saved anonymization details to: {json_output_path}")

    def _apply_replacements_to_docx(
        self,
        doc: Document,
        replacements: Dict[str, str],
        output_path: Path
    ) -> None:
        """
        Apply collected replacements directly to the DOCX document.

        This uses the explicit replacement map collected during processing,
        avoiding any diff-based guessing.
        """
        # Sort replacements by length (longest first) to avoid partial replacements
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

        # Apply replacements to paragraphs
        for para in doc.paragraphs:
            if para.text:
                new_text = para.text
                for original, replacement in sorted_replacements:
                    if original in new_text:
                        new_text = new_text.replace(original, replacement)

                if new_text != para.text:
                    self._replace_paragraph_text(para, new_text)

        # Apply replacements to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            new_text = para.text
                            for original, replacement in sorted_replacements:
                                if original in new_text:
                                    new_text = new_text.replace(original, replacement)

                            if new_text != para.text:
                                self._replace_paragraph_text(para, new_text)

        doc.save(output_path)

    def _replace_paragraph_text(self, paragraph, new_text: str) -> None:
        """
        Replace the text in a paragraph while trying to preserve formatting.

        This is a simplified approach that clears all runs and adds new text.
        """
        # If there's only one run or no runs, simple replacement
        if len(paragraph.runs) <= 1:
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            return

        # For multiple runs, try to preserve the first run's formatting
        # and put all text there
        first_run = paragraph.runs[0]

        # Clear all runs except the first
        for run in paragraph.runs[1:]:
            run.text = ""

        # Set the new text on the first run
        first_run.text = new_text

    def _extract_dates_with_regex(self, text: str) -> List[str]:
        """Extract all date/time patterns from text using regex."""
        patterns = [
            r'\b(\d{4}-\d{2}-\d{2}[T\s]\d{2}:\d{2}:\d{2})\b',
            r'\b(\d{4}-\d{2}-\d{2}\s+\d{1,2}:\d{2}\s*[AaPp][Mm])\b',
            r'\b(\d{4}-\d{2}-\d{2})\b',
            r'\b(\d{2}\.\d{2}\.\d{4})\b',
            r'\b(\d{2}/\d{2}/\d{4})\b',
            r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},?\s+\d{4})\b',
            r'\b(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})\b',
        ]

        found_dates: Set[str] = set()

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            found_dates.update(matches)

        return sorted(list(found_dates), key=len, reverse=True)

    def _phase1_shift_times(self, content: str) -> Tuple[str, List[DateTimeShift], Dict[str, str]]:
        """Phase 1: Find and shift all dates/times using regex extraction.

        Returns:
            Tuple of (modified content, list of shifts, dict of replacements)
        """
        modified_content = content
        all_shifts: List[DateTimeShift] = []
        replacements: Dict[str, str] = {}

        print("  Extracting dates using pattern matching...")
        found_dates = self._extract_dates_with_regex(content)
        print(f"  Found {len(found_dates)} unique date patterns to process")

        if not found_dates:
            return modified_content, all_shifts, replacements

        shifted_cache: Dict[str, str] = {}

        for date_str in found_dates:
            if date_str in shifted_cache:
                shifted_value = shifted_cache[date_str]
            else:
                try:
                    result = shift_datetime.invoke({
                        "datetime_str": date_str,
                        "offset_days": self.time_offset_days
                    })

                    if "[SHIFT_FAILED]" in result:
                        print(f"    Skip (invalid): {date_str}")
                        shifted_cache[date_str] = date_str
                        continue

                    shifted_value = result
                    shifted_cache[date_str] = shifted_value
                    print(f"    Shifted: {date_str} -> {shifted_value}")

                except Exception as e:
                    print(f"    Error shifting {date_str}: {e}")
                    shifted_cache[date_str] = date_str
                    continue

            if shifted_value != date_str:
                count = modified_content.count(date_str)
                if count > 0:
                    modified_content = modified_content.replace(date_str, shifted_value)
                    replacements[date_str] = shifted_value

                    all_shifts.append(DateTimeShift(
                        original_value=date_str,
                        shifted_value=shifted_value,
                        context=f"Found {count} occurrence(s)"
                    ))

        print(f"  Applied {len(all_shifts)} unique date shifts")
        return modified_content, all_shifts, replacements

    def _phase2_anonymize_pii(self, content: str) -> Tuple[str, int, Dict[str, str]]:
        """Phase 2: Anonymize all PII using agentic tool-calling with redact_text.

        Returns:
            Tuple of (modified content, number of redactions, dict of replacements)
        """
        modified_content = content
        total_redactions = 0
        all_replacements: Dict[str, str] = {}

        max_chunk_size = 8000

        if len(content) <= max_chunk_size:
            modified_content, total_redactions, all_replacements = self._anonymize_chunk_with_replacements(content, 0)
        else:
            chunks = self._split_into_chunks(content, max_chunk_size)
            print(f"  Processing {len(chunks)} chunks...")

            for chunk_num, (chunk_start, chunk_text) in enumerate(chunks):
                print(f"  Chunk {chunk_num + 1}/{len(chunks)} (chars {chunk_start}-{chunk_start + len(chunk_text)})")
                _, chunk_redactions, chunk_replacements = self._anonymize_chunk_with_replacements(chunk_text, chunk_start)
                total_redactions += chunk_redactions
                all_replacements.update(chunk_replacements)

            # Apply all collected redactions to the full content
            for original, redacted in all_replacements.items():
                modified_content = modified_content.replace(original, redacted)

        return modified_content, total_redactions, all_replacements

    def _split_into_chunks(self, content: str, max_size: int) -> List[Tuple[int, str]]:
        """Split content into overlapping chunks for processing."""
        chunks = []
        overlap = 500

        start = 0
        while start < len(content):
            end = min(start + max_size, len(content))

            if end < len(content):
                para_break = content.rfind('\n\n', start + max_size - 500, end)
                if para_break > start:
                    end = para_break
                else:
                    sent_break = content.rfind('. ', start + max_size - 200, end)
                    if sent_break > start:
                        end = sent_break + 1

            chunks.append((start, content[start:end]))
            start = end - overlap if end < len(content) else end

        return chunks

    def _anonymize_chunk_with_replacements(self, chunk: str, chunk_start: int) -> Tuple[str, int, Dict[str, str]]:
        """Anonymize a chunk of text using LLM with redact_text tool, returning replacements."""
        modified_chunk = chunk
        redactions = 0
        replacements: Dict[str, str] = {}

        prompt = self.prompt_config.get_text_anonymization_prompt(text_data=chunk)

        messages = [HumanMessage(content=prompt)]

        max_iterations = 50
        iteration = 0

        def invoke_llm_with_retry(msgs):
            return retry_with_backoff(
                lambda: self.llm_anonymize.invoke(msgs),
                config=self.retry_config,
                on_retry=create_retry_callback(prefix="    [LLM] "),
            )

        while iteration < max_iterations:
            iteration += 1

            try:
                response = invoke_llm_with_retry(messages)
                messages.append(response)

                if not response.tool_calls:
                    break

                for tool_call in response.tool_calls:
                    tool_name = tool_call["name"]
                    tool_args = tool_call["args"]

                    if tool_name == "redact_text":
                        text_to_redact = tool_args.get("text_to_redact", "")

                        result = redact_text.invoke(tool_args)

                        if "[REDACT_FAILED" not in result and text_to_redact:
                            if text_to_redact in modified_chunk:
                                modified_chunk = modified_chunk.replace(text_to_redact, result)
                                replacements[text_to_redact] = result
                                redactions += 1
                                display_text = text_to_redact[:40] + "..." if len(text_to_redact) > 40 else text_to_redact
                                print(f"    Redacted: '{display_text}'")

                        messages.append(ToolMessage(
                            content=f"Redacted: '{text_to_redact}' -> '{result}'",
                            tool_call_id=tool_call["id"]
                        ))

            except Exception as e:
                print(f"    Error during anonymization: {e}")
                break

        return modified_chunk, redactions, replacements

    def _phase3_verify_and_fix(
        self,
        original_content: str,
        anonymized_content: str
    ) -> Tuple[str, int, Dict[str, str]]:
        """Phase 3: Verification agent checks the anonymized output and fixes any issues.

        Returns:
            Tuple (modified content, number of fixes, dict of replacements)
        """
        modified_content = anonymized_content
        total_fixes = 0
        all_replacements: Dict[str, str] = {}

        max_chunk_size = 4000

        if len(original_content) <= max_chunk_size:
            modified_content, total_fixes, all_replacements = self._verify_chunk(
                original_content, anonymized_content, 0
            )
        else:
            orig_chunks = self._split_into_chunks(original_content, max_chunk_size)
            anon_chunks = self._split_into_chunks(anonymized_content, max_chunk_size)

            num_chunks = min(len(orig_chunks), len(anon_chunks))
            print(f"  Verifying {num_chunks} chunks...")

            for chunk_num in range(num_chunks):
                _, orig_chunk = orig_chunks[chunk_num]
                _, anon_chunk = anon_chunks[chunk_num]

                print(f"  Verifying chunk {chunk_num + 1}/{num_chunks}")
                _, chunk_fixes, chunk_replacements = self._verify_chunk(orig_chunk, anon_chunk, chunk_num)
                total_fixes += chunk_fixes
                all_replacements.update(chunk_replacements)

            # Apply all collected fixes to the full content
            for original, replacement in all_replacements.items():
                modified_content = modified_content.replace(original, replacement)

        return modified_content, total_fixes, all_replacements

    def _verify_chunk(
        self,
        original_chunk: str,
        anonymized_chunk: str,
        chunk_num: int
    ) -> Tuple[str, int, Dict[str, str]]:
        """Verify and fix a chunk of text.

        Returns:
            Tuple of (modified chunk, number of fixes, dict of replacements)
        """
        modified_chunk = anonymized_chunk
        fixes = 0
        replacements: Dict[str, str] = {}

        prompt = self.prompt_config.get_text_verification_prompt(
            original_text=original_chunk,
            anonymized_text=anonymized_chunk,
            time_offset=self.time_offset_days
        )

        messages = [HumanMessage(content=prompt)]

        max_iterations = 30
        iteration = 0

        def invoke_verify_with_retry(msgs):
            return retry_with_backoff(
                lambda: self.llm_verify.invoke(msgs),
                config=self.retry_config,
                on_retry=create_retry_callback(prefix="    [Verify] "),
            )

        while iteration < max_iterations:
            iteration += 1

            try:
                response = invoke_verify_with_retry(messages)
                messages.append(response)

                if not response.tool_calls:
                    break

                for tool_call in response.tool_calls:
                    tool_name = tool_call["name"]
                    tool_args = tool_call["args"]

                    if tool_name == "shift_datetime":
                        original_date = tool_args.get("datetime_str", "")
                        result = shift_datetime.invoke(tool_args)

                        if "[SHIFT_FAILED]" not in result:
                            if original_date in modified_chunk:
                                modified_chunk = modified_chunk.replace(original_date, result)
                                replacements[original_date] = result
                                fixes += 1
                                print(f"    Fixed date: {original_date} -> {result}")

                        messages.append(ToolMessage(
                            content=f"Date shifted: {original_date} -> {result}",
                            tool_call_id=tool_call["id"]
                        ))

                    elif tool_name == "redact_text":
                        text_to_redact = tool_args.get("text_to_redact", "")
                        result = redact_text.invoke(tool_args)

                        if "[REDACT_FAILED" not in result:
                            if text_to_redact in modified_chunk:
                                modified_chunk = modified_chunk.replace(text_to_redact, result)
                                replacements[text_to_redact] = result
                                fixes += 1
                                print(f"    Fixed PII: '{text_to_redact}' -> '{result}'")

                        messages.append(ToolMessage(
                            content=f"Text redacted: '{text_to_redact}' -> '{result}'",
                            tool_call_id=tool_call["id"]
                        ))

            except Exception as e:
                print(f"    Verification error: {e}")
                break

        return modified_chunk, fixes, replacements

    def _save_json_output(
        self,
        dates_shifted: List[DateTimeShift],
        pii_redactions_count: int,
        input_path: Path,
        output_path: Path,
        json_output_path: Path
    ) -> None:
        """Save anonymization details as JSON."""
        output_data = {
            "metadata": {
                "input_file": str(input_path.name),
                "output_file": str(output_path.name),
                "timestamp": datetime.now().isoformat(),
                "processing_method": "docx_anonymization",
                "time_offset_days": self.time_offset_days,
                "total_dates_shifted": len(dates_shifted),
                "total_pii_redactions": pii_redactions_count
            },
            "phase1_time_shifts": [
                {
                    "original": d.original_value,
                    "shifted": d.shifted_value,
                    "context": d.context
                }
                for d in dates_shifted
            ]
        }

        with open(json_output_path, "w", encoding="utf-8") as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
